#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档解析器 - 任务1
支持多种文件格式的解析
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any


class DocumentParser:
    """文档解析器，支持多种文件格式"""

    def __init__(self):
        """初始化解析器"""
        self.supported_formats = ['.txt', '.pdf', '.docx', '.md', '.json', '.csv']

    def parse_document(self, file_path: str) -> str:
        """
        解析文档并返回文本内容

        Args:
            file_path: 文件路径

        Returns:
            解析后的文本内容
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_extension = file_path.suffix.lower()

        if file_extension == '.txt':
            return self._parse_txt(file_path)
        elif file_extension == '.pdf':
            return self._parse_pdf(file_path)
        elif file_extension == '.docx':
            return self._parse_docx(file_path)
        elif file_extension == '.md':
            return self._parse_markdown(file_path)
        elif file_extension == '.json':
            return self._parse_json(file_path)
        elif file_extension == '.csv':
            return self._parse_csv(file_path)
        else:
            # 尝试作为纯文本读取
            try:
                return self._parse_txt(file_path)
            except:
                raise ValueError(f"不支持的文件格式: {file_extension}")

    def _parse_txt(self, file_path: Path) -> str:
        """解析TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except UnicodeDecodeError:
                # 最后尝试latin-1
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()

    def _parse_pdf(self, file_path: Path) -> str:
        """解析PDF文件"""
        try:
            # 尝试使用pdfplumber
            import pdfplumber

            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            return '\n\n'.join(text_content)

        except ImportError:
            # 如果pdfplumber不可用，尝试PyPDF2
            try:
                import PyPDF2

                text_content = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)

                return '\n\n'.join(text_content)

            except ImportError:
                return f"PDF文件: {file_path.name} (需要安装pdfplumber或PyPDF2库)\n安装命令: pip install pdfplumber"

        except Exception as e:
            raise ValueError(f"PDF解析失败: {e}")

    def _parse_docx(self, file_path: Path) -> str:
        """解析DOCX文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_content = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            return '\n\n'.join(text_content)

        except ImportError:
            return f"DOCX文件: {file_path.name} (需要安装python-docx库)\n安装命令: pip install python-docx"

        except Exception as e:
            raise ValueError(f"DOCX解析失败: {e}")

    def _parse_markdown(self, file_path: Path) -> str:
        """解析Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 简单清理markdown标记
        import re

        # 移除代码块
        content = re.sub(r'```[\s\S]*?```', '', content)

        # 移除行内代码
        content = re.sub(r'`([^`]+)`', r'\1', content)

        # 移除标题标记
        content = re.sub(r'#+\s*', '', content)

        # 移除粗体标记
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
        content = re.sub(r'__(.*?)__', r'\1', content)

        # 移除斜体标记
        content = re.sub(r'\*(.*?)\*', r'\1', content)
        content = re.sub(r'_(.*?)_', r'\1', content)

        # 移除链接，保留文本
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)

        # 移除图片
        content = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'[图片: \1]', content)

        return content.strip()

    def _parse_json(self, file_path: Path) -> str:
        """解析JSON文件"""
        import json

        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # 将JSON转换为可读文本
        return json.dumps(data, ensure_ascii=False, indent=2)

    def _parse_csv(self, file_path: Path) -> str:
        """解析CSV文件"""
        import csv

        content_lines = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    content_lines.append('\t'.join(row))
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, 'r', encoding='gbk') as f:
                reader = csv.reader(f)
                for row in reader:
                    content_lines.append('\t'.join(row))

        return '\n'.join(content_lines)

    def parse_image(self, file_path: str) -> str:
        """
        解析图片文件（OCR）

        Args:
            file_path: 图片路径

        Returns:
            识别的文本内容
        """
        try:
            from PIL import Image
            import pytesseract

            # 打开图片
            image = Image.open(file_path)

            # 进行OCR识别
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')

            return text.strip()

        except ImportError:
            return f"图片文件: {Path(file_path).name} (需要安装PIL和pytesseract库)\n安装命令: pip install Pillow pytesseract"

        except Exception as e:
            raise ValueError(f"图片解析失败: {e}")

    def get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取文件元数据"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        stat = file_path.stat()

        return {
            'filename': file_path.name,
            'path': str(file_path.absolute()),
            'size': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'extension': file_path.suffix.lower(),
            'created_time': stat.st_ctime,
            'modified_time': stat.st_mtime,
            'is_supported': file_path.suffix.lower() in self.supported_formats
        }

    def is_supported(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        return Path(file_path).suffix.lower() in self.supported_formats

    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式列表"""
        return self.supported_formats.copy()